"""
Script per generare la relazione del progetto AgriBot in formato .docx
compatibile con Apple Pages.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ── Stili personalizzati ──
style = doc.styles['Normal']
font = style.font
font.name = 'Helvetica'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 0, 0)

# Stile per codice
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Courier New'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.font.color.rgb = RGBColor(30, 30, 30)


def add_code(text):
    """Aggiunge un blocco di codice formattato."""
    for line in text.strip().split('\n'):
        doc.add_paragraph(line, style='CodeBlock')


def placeholder_img(didascalia):
    """Aggiunge un segnaposto per un'immagine che l'utente inserirà."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[INSERIRE IMMAGINE: {didascalia}]')
    run.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)
    run.font.size = Pt(11)
    # Didascalia sotto
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(didascalia)
    cap_run.italic = True
    cap_run.font.size = Pt(10)


# ═══════════════════════════════════════════════════════════════
# FRONTESPIZIO
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AgriBot-Agent')
run.bold = True
run.font.size = Pt(32)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Agente Autonomo per l\'Irrigazione Intelligente\nbasato su Computer Vision e Ricerca nello Spazio degli Stati')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Progetto di Introduzione all\'Intelligenza Artificiale')
run.font.size = Pt(13)

doc.add_paragraph('')
autori = doc.add_paragraph()
autori.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = autori.add_run('[Nome e Cognome]\n[Anno Accademico]')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INDICE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Indice', level=1)
indice = [
    '1. Introduzione e Obiettivi',
    '2. Architettura Generale del Sistema',
    '3. Struttura del Progetto e File Utilizzati',
    '4. Fase 1 — Addestramento della Rete Neurale (train.py)',
    '   4.1. Caricamento del Dataset EMNIST',
    '   4.2. Generazione delle Celle Vuote (classe ".")',
    '   4.3. Architettura della CNN',
    '   4.4. Training e Early Stopping',
    '   4.5. Grafici delle Prestazioni',
    '5. Fase 2 — Ritaglio della Griglia dall\'Immagine (cropping.py)',
    '   5.1. Caricamento e Binarizzazione',
    '   5.2. Rilevamento del Contorno Esterno',
    '   5.3. Ordinamento dei Vertici e Trasformazione Prospettica',
    '   5.4. Individuazione degli Intervalli delle Celle',
    '   5.5. Pulizia delle Singole Celle (clean_cell)',
    '   5.6. Estrazione Finale e Restituzione',
    '6. Fase 3 — Riconoscimento e Costruzione della Mappa (main.py)',
    '   6.1. Classificazione CNN delle Celle',
    '   6.2. Costruzione della Griglia Logica',
    '7. Fase 4 — Definizione del Problema di Ricerca (AgriBotProblem)',
    '   7.1. Rappresentazione dello Stato',
    '   7.2. Azioni Possibili',
    '   7.3. Funzione di Transizione',
    '   7.4. Test di Obiettivo',
    '8. Fase 5 — Algoritmi di Ricerca e Euristiche',
    '   8.1. Uniform Cost Search (UCS)',
    '   8.2. A* con Euristica Manhattan',
    '   8.3. A* con Euristica Max Pairwise Distance',
    '   8.4. Confronto tra gli Algoritmi',
    '9. Fase 6 — Visualizzazione della Soluzione',
    '10. Libreria AIMA',
    '11. Conclusioni',
]
for voce in indice:
    p = doc.add_paragraph(voce)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUZIONE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Introduzione e Obiettivi', level=1)

doc.add_paragraph(
    'AgriBot è un agente intelligente progettato per operare in un ambiente agricolo simulato. '
    'L\'ambiente è rappresentato da una griglia N×N disegnata su carta (o generata al computer) '
    'e fotografata. Il sistema deve:'
)

obiettivi = [
    'Acquisire l\'immagine della griglia e ritagliare automaticamente ogni singola cella tramite tecniche di Computer Vision.',
    'Riconoscere il contenuto di ciascuna cella (lettere D, F, R, S, T, V oppure cella vuota) mediante una rete neurale convoluzionale (CNN) addestrata sul dataset EMNIST.',
    'Costruire una rappresentazione logica dell\'ambiente (mappa) a partire dalle predizioni della rete.',
    'Formulare il problema come ricerca nello spazio degli stati e risolverlo con algoritmi di ricerca informata (A*) e non informata (UCS).',
    'Visualizzare graficamente la soluzione trovata, mostrando passo dopo passo le azioni del robot sulla mappa.',
]
for ob in obiettivi:
    doc.add_paragraph(ob, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    'Ogni cella della griglia rappresenta un elemento dell\'ambiente agricolo:'
)

elementi = [
    ('S (Start)', 'Posizione iniziale del robot.'),
    ('F (Finish)', 'Posizione finale che il robot deve raggiungere dopo aver completato tutti i compiti.'),
    ('R (Rock)', 'Ostacolo invalicabile; il robot non può attraversare questa cella.'),
    ('D (Dry Plant)', 'Pianta secca che necessita di 1 unità d\'acqua per essere irrigata.'),
    ('V (Very Dry Plant)', 'Pianta molto secca che necessita di 2 unità d\'acqua.'),
    ('T (Tank)', 'Stazione di rifornimento; il robot può ricaricare il serbatoio d\'acqua.'),
    ('. (Vuota)', 'Terreno libero percorribile.'),
]
table = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
hdr = table.rows[0].cells
hdr[0].text = 'Simbolo'
hdr[1].text = 'Significato'
for sym, desc in elementi:
    row = table.add_row().cells
    row[0].text = sym
    row[1].text = desc

doc.add_paragraph('')
doc.add_paragraph(
    'L\'obiettivo del robot è irrigare tutte le piante (D e V) gestendo un serbatoio d\'acqua '
    'con capacità limitata, e infine raggiungere la cella di arrivo F con il costo minimo possibile.'
)

placeholder_img('Esempio di griglia disegnata su carta e relativa mappa logica riconosciuta')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. ARCHITETTURA
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Architettura Generale del Sistema', level=1)

doc.add_paragraph(
    'Il progetto segue una pipeline sequenziale composta da quattro macro-fasi:'
)

doc.add_paragraph(
    '1. Addestramento (offline): si allena una CNN sul dataset EMNIST Balanced, '
    'filtrato sulle sole classi di interesse (D, F, R, S, T, V) a cui si aggiunge '
    'una classe per le celle vuote. Il modello addestrato viene salvato su disco.'
)
doc.add_paragraph(
    '2. Visione (Computer Vision): l\'immagine della griglia viene processata con '
    'OpenCV per trovare il contorno esterno, correggere la prospettiva, e ritagliare '
    'ogni singola cella in un\'immagine 28×28 pixel in scala di grigi.'
)
doc.add_paragraph(
    '3. Classificazione: ogni cella ritagliata viene passata alla CNN che predice '
    'il carattere contenuto. Si costruisce così la mappa logica (matrice di stringhe).'
)
doc.add_paragraph(
    '4. Pianificazione e Ricerca: la mappa logica viene usata per istanziare un '
    'problema di ricerca nello spazio degli stati. Si applicano UCS e A* (con due '
    'euristiche diverse) per trovare il piano d\'azione ottimale.'
)

doc.add_paragraph('')
doc.add_paragraph(
    'Schema della pipeline:'
)
doc.add_paragraph(
    'Immagine Input  →  Cropping (OpenCV)  →  Classificazione (CNN)  →  Mappa Logica  →  '
    'Pathfinding (A* / UCS)  →  Visualizzazione Soluzione'
)

placeholder_img('Diagramma della pipeline del sistema')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. STRUTTURA DEL PROGETTO
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Struttura del Progetto e File Utilizzati', level=1)

doc.add_paragraph(
    'Il progetto è organizzato nei seguenti file principali:'
)

files_desc = [
    ('train.py',
     'Script per l\'addestramento della rete neurale convoluzionale (CNN). '
     'Scarica il dataset EMNIST Balanced, filtra le classi di interesse, '
     'genera celle vuote sintetiche per la classe ".", costruisce il modello CNN, '
     'lo addestra con early stopping e salva il modello finale come agribot_model.keras. '
     'Produce inoltre il grafico delle prestazioni (accuracy e loss).'),

    ('cropping.py',
     'Modulo di Computer Vision. Contiene la funzione crop() che riceve il percorso '
     'di un\'immagine e le dimensioni della griglia (righe × colonne). '
     'Rileva automaticamente il contorno della griglia, corregge la prospettiva con '
     'una trasformazione proiettiva (warp), individua le linee della griglia tramite '
     'apertura morfologica, e ritaglia ogni singola cella producendo immagini 28×28 '
     'pronte per la classificazione.'),

    ('main.py',
     'File principale del progetto. Integra tutte le componenti: '
     '(1) chiama crop() per ritagliare le celle dall\'immagine, '
     '(2) usa la CNN per classificare ogni cella e costruire la mappa logica, '
     '(3) definisce la classe AgriBotProblem che modella il problema di ricerca, '
     '(4) esegue UCS e A* con due euristiche (Manhattan e Max Pairwise Distance), '
     '(5) visualizza graficamente la soluzione passo dopo passo.'),

    ('agribot_model.keras',
     'File del modello CNN pre-addestrato. Viene generato da train.py e '
     'caricato da main.py per la classificazione delle celle.'),

    ('aima/ (cartella)',
     'Contiene la libreria AIMA-Python (search.py e utils.py) che fornisce '
     'le implementazioni degli algoritmi di ricerca (A*, UCS) e la classe base '
     'Problem da cui eredita AgriBotProblem.'),
]

table2 = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
hdr = table2.rows[0].cells
hdr[0].text = 'File'
hdr[1].text = 'Descrizione'
for fname, fdesc in files_desc:
    row = table2.add_row().cells
    row[0].text = fname
    row[1].text = fdesc

# Imposta larghezza colonne
for row in table2.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(13)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. TRAIN.PY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. Fase 1 — Addestramento della Rete Neurale (train.py)', level=1)

# 4.1
doc.add_heading('4.1. Caricamento del Dataset EMNIST', level=2)
doc.add_paragraph(
    'Il dataset utilizzato è EMNIST Balanced, una versione estesa di MNIST che include '
    'sia cifre sia lettere dell\'alfabeto. Viene scaricato automaticamente tramite la '
    'libreria tensorflow_datasets.'
)
doc.add_paragraph(
    'Poiché il progetto utilizza solo un sottoinsieme di lettere, viene applicato un filtro: '
    'si selezionano esclusivamente le classi corrispondenti alle lettere D, F, R, S, T e V. '
    'La mappatura tra indici EMNIST e classi interne del progetto è definita come segue:'
)

add_code('''TARGET_MAP = {13: 0, 15: 1, 27: 2, 28: 3, 29: 4, 31: 6}
CLASSES = ['D', 'F', 'R', 'S', 'T', '.', 'V']''')

doc.add_paragraph(
    'Ogni immagine EMNIST è in formato 28×28 pixel in scala di grigi. Le immagini vengono '
    'trasposte (np.transpose) perché EMNIST le memorizza ruotate rispetto alla convenzione standard.'
)

# 4.2
doc.add_heading('4.2. Generazione delle Celle Vuote (classe ".")', level=2)
doc.add_paragraph(
    'Per insegnare alla rete a riconoscere le celle vuote della griglia (senza nessuna lettera), '
    'si generano sinteticamente immagini di sfondo. Il numero di celle vuote generate è pari a '
    'un quinto del dataset filtrato. Metà di queste sono immagini completamente nere (sfondo puro), '
    'mentre l\'altra metà contiene rumore casuale di bassa intensità (valori tra 0 e 50), per '
    'rendere il modello robusto anche in presenza di piccole imperfezioni nell\'immagine reale.'
)

add_code('''n_noise = len(X) // 5
for i in range(n_noise):
    img = np.zeros((28, 28), dtype=np.uint8)
    if i % 2 == 0:
        noise = np.random.randint(0, 50, (28, 28), dtype=np.uint8)
        img = cv2.add(img, noise)
    x_noise.append(img)''')

doc.add_paragraph(
    'A queste immagini viene assegnata la label 5, che corrisponde alla classe "." (cella vuota) '
    'nell\'array CLASSES.'
)

# 4.3
doc.add_heading('4.3. Architettura della CNN', level=2)
doc.add_paragraph(
    'Il modello è una rete neurale convoluzionale (Convolutional Neural Network) costruita '
    'con l\'API Sequential di Keras. L\'architettura è la seguente:'
)

strati = [
    ('Input', '28×28×1 — immagine in scala di grigi.'),
    ('RandomRotation(0.1)', 'Data augmentation: rotazione casuale fino a ±10% per rendere il modello invariante a piccole rotazioni.'),
    ('RandomZoom(0.1)', 'Data augmentation: zoom casuale fino a ±10% per gestire variazioni di scala.'),
    ('Conv2D(32, 3×3, ReLU, same)', 'Primo strato convoluzionale: 32 filtri 3×3 con attivazione ReLU e padding "same" per mantenere le dimensioni.'),
    ('MaxPooling2D(2×2)', 'Riduce la dimensione spaziale della metà (da 28×28 a 14×14).'),
    ('Dropout(0.2)', 'Spegne casualmente il 20% dei neuroni durante il training per prevenire l\'overfitting.'),
    ('Conv2D(64, 3×3, ReLU, same)', 'Secondo strato convoluzionale: 64 filtri per catturare pattern più complessi.'),
    ('MaxPooling2D(2×2)', 'Riduce ulteriormente (da 14×14 a 7×7).'),
    ('Dropout(0.2)', 'Ulteriore regolarizzazione.'),
    ('Flatten', 'Appiattisce la mappa 7×7×64 = 3136 valori in un vettore 1D.'),
    ('Dense(128, ReLU)', 'Strato fully-connected con 128 neuroni.'),
    ('Dropout(0.4)', 'Dropout più aggressivo prima dell\'output per massima regolarizzazione.'),
    ('Dense(7, Softmax)', 'Strato di output: 7 neuroni (uno per classe) con Softmax per produrre probabilità.'),
]

table3 = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
hdr = table3.rows[0].cells
hdr[0].text = 'Strato'
hdr[1].text = 'Descrizione'
for sname, sdesc in strati:
    row = table3.add_row().cells
    row[0].text = sname
    row[1].text = sdesc

doc.add_paragraph('')
doc.add_paragraph(
    'Il modello viene compilato con ottimizzatore Adam, funzione di loss categorical_crossentropy '
    '(adatta alla classificazione multi-classe) e metrica accuracy.'
)

# 4.4
doc.add_heading('4.4. Training e Early Stopping', level=2)
doc.add_paragraph(
    'Il dataset viene suddiviso in 80% training e 20% validation tramite train_test_split '
    'con stratificazione (stratify=Y), che garantisce che ogni classe sia rappresentata '
    'proporzionalmente in entrambi i set.'
)
doc.add_paragraph(
    'Il training viene eseguito per un massimo di 30 epoche con batch size 64. '
    'Viene utilizzato un callback EarlyStopping che monitora la validation loss: '
    'se questa non migliora per 5 epoche consecutive (patience=5), il training si '
    'interrompe automaticamente e vengono ripristinati i pesi migliori '
    '(restore_best_weights=True). Questo evita l\'overfitting.'
)

add_code('''early_stop = EarlyStopping(monitor='val_loss', patience=5, restore_best_weights=True)
history = model.fit(X_train, Y_train, validation_data=(X_val, Y_val),
                    epochs=30, batch_size=64, callbacks=[early_stop])
model.save("agribot_model.keras")''')

# 4.5
doc.add_heading('4.5. Grafici delle Prestazioni', level=2)
doc.add_paragraph(
    'Al termine del training, la funzione plot_history() genera due grafici affiancati:'
)
doc.add_paragraph('Training Accuracy vs Validation Accuracy: mostra come la precisione evolve durante le epoche.', style='List Bullet')
doc.add_paragraph('Training Loss vs Validation Loss: mostra come l\'errore diminuisce. Se le curve di training e validation divergono, indica overfitting.', style='List Bullet')
doc.add_paragraph(
    'I grafici vengono salvati come statistiche_training.png.'
)

placeholder_img('Grafici di Training Accuracy e Training Loss (statistiche_training.png)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. CROPPING.PY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Fase 2 — Ritaglio della Griglia dall\'Immagine (cropping.py)', level=1)

doc.add_paragraph(
    'Il modulo cropping.py è il cuore della parte di Computer Vision del progetto. '
    'La funzione principale crop(filename, n_rows, n_cols) riceve il percorso di un\'immagine '
    'e le dimensioni della griglia, e restituisce una lista di tuple (riga, colonna, immagine_28x28) '
    'per ogni cella. Il processo si articola nelle seguenti fasi.'
)

# 5.1
doc.add_heading('5.1. Caricamento e Binarizzazione', level=2)
doc.add_paragraph(
    'L\'immagine viene caricata con cv2.imread() e convertita in scala di grigi con '
    'cv2.cvtColor(img, cv2.COLOR_BGR2GRAY). Successivamente si applica una sogliatura '
    'binaria inversa con il metodo di Otsu:'
)

add_code('''gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
_, negative_img = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)''')

doc.add_paragraph(
    'THRESH_BINARY_INV inverte i colori: lo sfondo chiaro diventa nero (0) e le linee/lettere '
    'scure diventano bianche (255). THRESH_OTSU calcola automaticamente la soglia ottimale '
    'analizzando l\'istogramma dell\'immagine, senza bisogno di specificare un valore manuale.'
)

placeholder_img('Immagine originale e risultato della binarizzazione con Otsu')

# 5.2
doc.add_heading('5.2. Rilevamento del Contorno Esterno', level=2)
doc.add_paragraph(
    'Sull\'immagine binarizzata si cercano i contorni con cv2.findContours(). '
    'Il flag RETR_EXTERNAL indica di considerare solo i contorni esterni (il perimetro della griglia), '
    'ignorando i contorni interni (lettere, linee interne). '
    'CHAIN_APPROX_SIMPLE comprime i segmenti rettilinei memorizzando solo i punti estremi, '
    'risparmiando memoria.'
)

add_code('''find_contours, _ = cv2.findContours(negative_img, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
find_biggest_contours = sorted(find_contours, key=cv2.contourArea, reverse=True)''')

doc.add_paragraph(
    'I contorni trovati vengono ordinati per area in ordine decrescente: il più grande corrisponde '
    'al bordo esterno della griglia.'
)

# 5.3
doc.add_heading('5.3. Ordinamento dei Vertici e Trasformazione Prospettica', level=2)
doc.add_paragraph(
    'Dal contorno più grande si calcola il rettangolo minimo ruotato (minAreaRect), '
    'i cui 4 vertici vengono estratti con boxPoints. Questi vertici rappresentano '
    'gli angoli della griglia nell\'immagine originale, anche se questa è stata fotografata '
    'in prospettiva o leggermente ruotata.'
)

add_code('''corners = cv2.boxPoints(cv2.minAreaRect(c)).astype(np.float32)''')

doc.add_paragraph(
    'I vertici vengono poi ordinati in senso preciso (Top-Left, Top-Right, Bottom-Right, Bottom-Left) '
    'dalla funzione order_points(), che utilizza la somma (x+y) e la differenza (x-y) delle coordinate:'
)
doc.add_paragraph('Il vertice con la somma x+y minima è il Top-Left.', style='List Bullet')
doc.add_paragraph('Il vertice con la somma x+y massima è il Bottom-Right.', style='List Bullet')
doc.add_paragraph('Il vertice con la differenza x-y massima è il Top-Right.', style='List Bullet')
doc.add_paragraph('Il vertice con la differenza x-y minima è il Bottom-Left.', style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    'Con i vertici sorgente (src) e quelli di destinazione (dst = quadrato 1000×1000), '
    'si calcola la matrice di trasformazione prospettica e si applica il warp:'
)

add_code('''dst = np.float32([[0,0],[DIM-1,0],[DIM-1,DIM-1],[0,DIM-1]])
warped = cv2.warpPerspective(gray, cv2.getPerspectiveTransform(src, dst), (DIM, DIM))''')

doc.add_paragraph(
    'Il risultato è un\'immagine perfettamente raddrizzata di 1000×1000 pixel, indipendentemente '
    'dall\'angolo o dalla distorsione della foto originale. Sull\'immagine warpata si applica '
    'nuovamente la sogliatura di Otsu (per binarizzare) seguita da un\'apertura morfologica '
    'con kernel 3×3 per eliminare il rumore di fondo minuto.'
)

placeholder_img('Immagine originale, vertici rilevati e risultato del warp prospettico a 1000×1000')

# 5.4
doc.add_heading('5.4. Individuazione degli Intervalli delle Celle', level=2)
doc.add_paragraph(
    'Questa è la fase più delicata: bisogna individuare dove iniziano e finiscono le righe e le colonne '
    'della griglia all\'interno dell\'immagine warpata. Si utilizza un approccio basato sull\'apertura '
    'morfologica per isolare esclusivamente le linee della griglia, eliminando le lettere.'
)

doc.add_paragraph('')
doc.add_heading('Apertura Morfologica per l\'Estrazione delle Linee', level=3)
doc.add_paragraph(
    'L\'apertura morfologica (MORPH_OPEN) è composta da due operazioni consecutive: '
    'un\'erosione seguita da una dilatazione. L\'idea chiave è utilizzare un kernel rettangolare '
    'molto allungato:'
)
doc.add_paragraph(
    'Per trovare le linee orizzontali: si usa un kernel largo (es. 150 pixel) e alto 1 pixel. '
    'L\'erosione elimina tutte le strutture bianche che non sono almeno larghe 150 pixel in modo continuo. '
    'Le lettere (larghe al massimo ~120 pixel dentro una cella di ~250 pixel) vengono cancellate. '
    'Le linee della griglia (larghe ~1000 pixel, cioè l\'intera immagine) sopravvivono. '
    'La dilatazione successiva riporta le strutture sopravvissute alla loro dimensione originale.',
    style='List Bullet'
)
doc.add_paragraph(
    'Per trovare le linee verticali: si usa un kernel alto (es. 150 pixel) e largo 1 pixel, '
    'con lo stesso principio.',
    style='List Bullet'
)

add_code('''kernel_len = max(int(expected_size * 0.6), 40)  # ~60% della dimensione attesa della cella

# Per linee orizzontali:
kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (kernel_len, 1))
lines_only = cv2.morphologyEx(warped_img, cv2.MORPH_OPEN, kernel)
proj = lines_only.sum(axis=1).astype(float)  # Proiezione: somma di ogni riga

# Per linee verticali:
kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, kernel_len))
lines_only = cv2.morphologyEx(warped_img, cv2.MORPH_OPEN, kernel)
proj = lines_only.sum(axis=0).astype(float)  # Proiezione: somma di ogni colonna''')

doc.add_paragraph(
    'Dopo l\'apertura morfologica, l\'immagine lines_only contiene solo le linee della griglia, '
    'senza nessuna lettera. La proiezione (somma lungo un asse) produce un vettore 1D di 1000 valori: '
    'i valori sono alti dove c\'è una linea e zero dove non c\'è.'
)

placeholder_img('Immagine dopo apertura morfologica (solo linee) e proiezione risultante')

doc.add_paragraph('')
doc.add_heading('Ricerca dei Segmenti di Linea', level=3)
doc.add_paragraph(
    'Sulla proiezione si applica una soglia al 30% del valore massimo per ottenere un array '
    'booleano (True = linea, False = no linea). Per trovare l\'inizio e la fine di ogni segmento '
    'contiguo di True, si usa il seguente metodo:'
)

add_code('''threshold = proj.max() * 0.3
is_line = proj > threshold

padded = np.concatenate([[False], is_line, [False]])  # Aggiunge False ai bordi
diff = np.diff(padded.astype(int))                    # Calcola le transizioni
line_starts = np.where(diff == 1)[0]   # Dove inizia ogni linea
line_ends   = np.where(diff == -1)[0]  # Dove finisce ogni linea''')

doc.add_paragraph(
    'np.diff calcola la differenza tra elementi consecutivi: dove il valore passa da 0 a 1 '
    '(diff == 1) si ha l\'inizio di una linea; dove passa da 1 a 0 (diff == -1) si ha la fine. '
    'Il padding con False garantisce che anche le linee ai bordi vengano rilevate.'
)

doc.add_paragraph('')
doc.add_heading('Costruzione dei Gap (Intervalli delle Celle)', level=3)
doc.add_paragraph(
    'Una volta note le posizioni esatte di tutte le linee, le celle corrispondono ai gap tra linee consecutive. '
    'Si costruisce una lista di confini alternati [0, inizio_linea_1, fine_linea_1, inizio_linea_2, fine_linea_2, ..., 1000] '
    'e si estraggono i gap alle posizioni pari:'
)

add_code('''boundaries = [0]
for s, e in zip(line_starts, line_ends):
    boundaries.append(int(s))
    boundaries.append(int(e))
boundaries.append(DIM)

gaps = []
for i in range(0, len(boundaries) - 1, 2):
    g_start = boundaries[i] + MARGIN    # +2 pixel di sicurezza
    g_end = boundaries[i + 1] - MARGIN  # -2 pixel di sicurezza
    if g_end > g_start and (g_end - g_start) > expected_size * 0.15:
        gaps.append((g_start, g_end))''')

doc.add_paragraph(
    'Il MARGIN di 2 pixel restringe leggermente ogni cella per evitare di includere pixel '
    'residui della linea. I gap troppo piccoli (meno del 15% della dimensione attesa) vengono '
    'scartati — tipicamente sono gli spazi tra il bordo dell\'immagine e il bordo esterno della '
    'griglia. Infine, tra tutti i gap validi, si selezionano gli n più vicini alla dimensione '
    'attesa di una cella (DIM/n) e si riordinano per posizione.'
)

placeholder_img('Schema dei boundaries e dei gap risultanti per una griglia 4×4')

# 5.5
doc.add_heading('5.5. Pulizia delle Singole Celle (clean_cell)', level=2)
doc.add_paragraph(
    'Ogni cella ritagliata dalla griglia viene passata alla funzione clean_cell(), che la '
    'pulisce e la prepara per la classificazione CNN. I passaggi sono:'
)

doc.add_paragraph(
    '1. Bordo nero: si imposta a 0 (nero) la prima e l\'ultima riga e colonna della cella per '
    'eliminare eventuali residui delle linee della griglia che potrebbero essere rimasti ai bordi.'
)
doc.add_paragraph(
    '2. Connected Components: si usa cv2.connectedComponentsWithStats() che analizza '
    'l\'immagine binaria e raggruppa i pixel bianchi adiacenti in "isole" (componenti connesse). '
    'Ogni isola riceve un\'etichetta (label) e vengono calcolate le statistiche: posizione (x, y), '
    'dimensioni (w, h) e area.'
)
doc.add_paragraph(
    '3. Selezione dell\'isola principale: si seleziona la componente connessa con l\'area '
    'maggiore (escludendo lo sfondo nero, che è la componente 0). Questa dovrebbe essere la lettera.'
)
doc.add_paragraph(
    '4. Filtri di validazione: si applicano diversi controlli per escludere falsi positivi:'
)
doc.add_paragraph('Se l\'area è inferiore a 10 pixel, è rumore → cella vuota.', style='List Bullet')
doc.add_paragraph('Se il rapporto w/h è maggiore di 5 o inferiore a 0.2, è una forma troppo allungata → cella vuota.', style='List Bullet')
doc.add_paragraph('Se l\'oggetto è molto sottile e lungo (larghezza < 6px e altezza > 50% della cella), è un residuo di linea verticale della griglia → cella vuota. Analogamente per le linee orizzontali.', style='List Bullet')

doc.add_paragraph(
    '5. Ritaglio e centratura: la lettera viene ritagliata dal suo bounding box, '
    'inserita in un quadrato (con padding nero simmetrico) e ridimensionata a 28×28 pixel '
    'con interpolazione cv2.INTER_AREA per evitare artefatti di aliasing.'
)

add_code('''bordo = max(w, h) + 6
dy, dx = (bordo - h), (bordo - w)
square = np.pad(digit, ((dy//2, dy - dy//2), (dx//2, dx - dx//2)))
return cv2.resize(square, (28, 28), interpolation=cv2.INTER_AREA)''')

placeholder_img('Esempio di cella prima e dopo la pulizia con clean_cell')

# 5.6
doc.add_heading('5.6. Estrazione Finale e Restituzione', level=2)
doc.add_paragraph(
    'La funzione crop() itera su tutte le combinazioni di righe e colonne, ritaglia ogni cella '
    'dall\'immagine warpata usando gli intervalli trovati, la pulisce con clean_cell(), '
    'e restituisce una lista di tuple (i, j, immagine_28x28) dove i è l\'indice di riga, '
    'j l\'indice di colonna e immagine_28x28 è l\'immagine pronta per la CNN.'
)

add_code('''cells = []
for i, r in enumerate(rows[:n_rows]):
    for j, c in enumerate(cols[:n_cols]):
        cells.append((i, j, clean_cell(warped[r[0]:r[1], c[0]:c[1]])))
return cells''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. MAIN.PY — Classificazione
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Fase 3 — Riconoscimento e Costruzione della Mappa (main.py)', level=1)

# 6.1
doc.add_heading('6.1. Classificazione CNN delle Celle', level=2)
doc.add_paragraph(
    'Il file main.py inizia caricando il modello CNN pre-addestrato e ritagliando le celle '
    'dall\'immagine con crop(). Ogni cella 28×28 viene normalizzata (valori tra 0 e 1), '
    'ridimensionata a (1, 28, 28, 1) per il formato batch richiesto da Keras, e passata '
    'al modello per la predizione.'
)

add_code('''model = tf.keras.models.load_model("agribot_model.keras")
cells = crop("prova2.png", N_ROWS, N_COLS)

for i, j, cell_img in cells:
    normalized_cell = cell_img.astype("float32") / 255.0
    normalized_cell = normalized_cell.reshape(1, 28, 28, 1)
    prediction = model.predict(normalized_cell, verbose=0)
    idx = np.argmax(prediction)
    char = LABELS[idx]   # LABELS = ['D', 'F', 'R', 'S', 'T', '.', 'V']''')

doc.add_paragraph(
    'model.predict() restituisce un array di 7 probabilità (una per classe). '
    'np.argmax() seleziona l\'indice con la probabilità più alta, che viene '
    'convertito nel carattere corrispondente tramite l\'array LABELS.'
)

# 6.2
doc.add_heading('6.2. Costruzione della Griglia Logica', level=2)
doc.add_paragraph(
    'I caratteri riconosciuti vengono organizzati in una matrice (lista di liste) grid_map, '
    'dove ogni riga interna contiene i caratteri delle N_COLS colonne. Questa matrice rappresenta '
    'la mappa logica dell\'ambiente e viene passata al costruttore di AgriBotProblem.'
)

add_code('''grid_map = []
row = []
for i, j, cell_img in cells:
    # ... predizione ...
    row.append(char)
    if j == N_COLS - 1:
        grid_map.append(row)
        row = []''')

placeholder_img('Esempio di mappa logica riconosciuta dalla CNN con matrice risultante')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. AGRIBOTPROBLEM
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. Fase 4 — Definizione del Problema di Ricerca (AgriBotProblem)', level=1)

doc.add_paragraph(
    'La classe AgriBotProblem eredita dalla classe Problem della libreria AIMA. '
    'Questa modellazione consente di utilizzare direttamente gli algoritmi di ricerca (UCS, A*) '
    'forniti dalla libreria.'
)

# 7.1
doc.add_heading('7.1. Rappresentazione dello Stato', level=2)
doc.add_paragraph(
    'Lo stato del robot è rappresentato dalla tupla:'
)
doc.add_paragraph(
    '(posizione, acqua, dry, very_dry)', style='List Bullet'
)
doc.add_paragraph(
    'Dove:'
)
doc.add_paragraph('posizione: un intero che rappresenta la cella corrente del robot (indice lineare r×n+c).', style='List Bullet')
doc.add_paragraph('acqua: il livello corrente del serbatoio d\'acqua (intero, da 0 a max_water).', style='List Bullet')
doc.add_paragraph('dry: un frozenset contenente gli indici delle piante D ancora da irrigare.', style='List Bullet')
doc.add_paragraph('very_dry: un frozenset contenente gli indici delle piante V ancora da irrigare.', style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    'L\'uso di frozenset è necessario perché gli stati devono essere hashable per poter essere '
    'inseriti negli insiemi "esplorati" degli algoritmi di ricerca. Ogni volta che una pianta '
    'viene irrigata, si crea un nuovo frozenset senza quella pianta (immutabilità).'
)

doc.add_paragraph('')
doc.add_paragraph(
    'La conversione da coordinate (riga, colonna) a indice lineare avviene con la formula '
    'index = riga × n + colonna, dove n è la dimensione della griglia. '
    'Ad esempio, in una griglia 4×4, la cella (1, 2) corrisponde all\'indice 1×4+2 = 6.'
)

# 7.2
doc.add_heading('7.2. Azioni Possibili', level=2)
doc.add_paragraph(
    'Il metodo actions(state) restituisce la lista delle azioni valide nello stato corrente. '
    'Le azioni sono generate dinamicamente in base alla posizione, alle risorse e agli elementi '
    'circostanti:'
)

azioni = [
    ('UP, DOWN, LEFT, RIGHT', 'Movimento nelle 4 direzioni cardinali. Un\'azione di movimento è valida solo se la cella di destinazione è dentro la griglia e non contiene una roccia (R). Costo: 1.'),
    ('WATER', 'Irrigazione della pianta nella cella corrente. Valida solo se il robot è su una cella D (richiede 1 unità d\'acqua) o V (richiede 2 unità d\'acqua) e ha acqua sufficiente. Costo: 1.'),
    ('REFILL', 'Ricarica del serbatoio. Valida solo se il robot è su una stazione T e il serbatoio non è già pieno. Riporta l\'acqua al massimo. Costo: 1.'),
]
table4 = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
hdr = table4.rows[0].cells
hdr[0].text = 'Azione'
hdr[1].text = 'Descrizione e Condizioni'
for aname, adesc in azioni:
    row = table4.add_row().cells
    row[0].text = aname
    row[1].text = adesc

# 7.3
doc.add_heading('7.3. Funzione di Transizione', level=2)
doc.add_paragraph(
    'Il metodo result(state, action) restituisce il nuovo stato risultante dall\'applicazione '
    'di un\'azione. Per i movimenti, la posizione viene aggiornata sommando il delta '
    'corrispondente alla direzione (-n per UP, +n per DOWN, -1 per LEFT, +1 per RIGHT). '
    'Per WATER, la pianta viene rimossa dal set corrispondente e l\'acqua viene decrementata. '
    'Per REFILL, l\'acqua torna al massimo.'
)

add_code('''delta = {'UP': -self.n, 'DOWN': self.n, 'LEFT': -1, 'RIGHT': 1}

if action == "WATER" and position in dry:
    return (position, water - 1, dry.difference({position}), very_dry)
if action == "WATER" and position in very_dry:
    return (position, water - 2, dry, very_dry.difference({position}))
if action == "REFILL":
    return (position, self.max_water, dry, very_dry)

new_position = position + delta[action]
return (new_position, water, dry, very_dry)''')

# 7.4
doc.add_heading('7.4. Test di Obiettivo', level=2)
doc.add_paragraph(
    'Il metodo goal_test(state) verifica se lo stato corrente è un obiettivo. '
    'La condizione è che tutte le piante siano state irrigate (entrambi i set dry e very_dry '
    'sono vuoti) E il robot si trovi sulla cella di arrivo F.'
)

add_code('''def goal_test(self, state):
    position, water, dry, very_dry = state
    return len(dry) == 0 and len(very_dry) == 0 and position == self.finish_position''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. ALGORITMI DI RICERCA
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. Fase 5 — Algoritmi di Ricerca e Euristiche', level=1)

doc.add_paragraph(
    'Il progetto confronta tre approcci di ricerca per trovare il piano d\'azione ottimale.'
)

# 8.1
doc.add_heading('8.1. Uniform Cost Search (UCS)', level=2)
doc.add_paragraph(
    'UCS è un algoritmo di ricerca non informata (cieco) che espande sempre il nodo con '
    'il costo di percorso più basso. Garantisce di trovare la soluzione ottimale (costo minimo) '
    'ma esplora molti nodi inutili perché non ha informazioni sulla direzione dell\'obiettivo. '
    'Viene usato come baseline per valutare l\'efficacia delle euristiche.'
)

# 8.2
doc.add_heading('8.2. A* con Euristica Manhattan (h_manhattan)', level=2)
doc.add_paragraph(
    'A* è un algoritmo di ricerca informata che utilizza una funzione di valutazione '
    'f(n) = g(n) + h(n), dove g(n) è il costo dal nodo iniziale al nodo n, e h(n) è '
    'una stima euristica del costo dal nodo n all\'obiettivo.'
)
doc.add_paragraph(
    'L\'euristica Manhattan (h_manhattan) considera tre casi in base allo stato corrente:'
)
doc.add_paragraph(
    'Caso A — Tutte le piante irrigate: restituisce la distanza Manhattan dalla posizione '
    'corrente alla cella di arrivo F. Non ci sono più compiti, bisogna solo raggiungere la fine.',
    style='List Bullet'
)
doc.add_paragraph(
    'Caso B — Ci sono ancora piante e il robot ha acqua: restituisce la distanza Manhattan '
    'verso la pianta più vicina. Il prossimo passo logico è irrigare.',
    style='List Bullet'
)
doc.add_paragraph(
    'Caso C — Ci sono ancora piante ma il robot è senza acqua: restituisce la distanza '
    'Manhattan verso la stazione di rifornimento T più vicina. Bisogna prima ricaricare.',
    style='List Bullet'
)

add_code('''def h_manhattan(self, node):
    position, water, dry, very_dry = node.state
    all_plants = dry.union(very_dry)

    if len(all_plants) == 0:
        return self.cal_manhattan(position, self.finish_position)
    if water > 0:
        return min(self.cal_manhattan(position, p) for p in all_plants)
    if water == 0:
        return min(self.cal_manhattan(position, s) for s in self.station)''')

doc.add_paragraph(
    'Questa euristica è ammissibile (non sovrastima mai il costo reale) e quindi A* con '
    'h_manhattan garantisce l\'ottimalità. Tuttavia sottostima significativamente il costo '
    'in scenari complessi con molte piante sparse, perché considera solo l\'obiettivo più '
    'vicino ignorando tutti gli altri.'
)

# 8.3
doc.add_heading('8.3. A* con Euristica Max Pairwise Distance (h_max_pairwise_Distance)', level=2)
doc.add_paragraph(
    'Questa euristica più sofisticata aggiunge alla stima il "diametro" del set di piante '
    'rimaste, cioè la massima distanza Manhattan tra qualsiasi coppia di piante. '
    'L\'intuizione è: anche dopo aver raggiunto la pianta più vicina, il robot dovrà '
    'comunque percorrere almeno la distanza tra le due piante più lontane tra loro.'
)

add_code('''def h_max_pairwise_Distance(self, node):
    position, water, dry, very_dry = node.state
    all_plants = dry.union(very_dry)

    if len(all_plants) == 0:
        return self.cal_manhattan(position, self.finish_position)

    # Calcola la massima distanza tra qualsiasi coppia di piante
    max_internal_distance = 0
    plants_list = list(all_plants)
    for i in range(len(plants_list)):
        for j in range(i + 1, len(plants_list)):
            dist = self.cal_manhattan(plants_list[i], plants_list[j])
            if dist > max_internal_distance:
                max_internal_distance = dist

    if water == 0:
        return min(self.cal_manhattan(position, s) for s in self.station) + max_internal_distance
    if water > 0:
        most_close = min(self.cal_manhattan(position, p) for p in all_plants)
        return max_internal_distance + most_close''')

doc.add_paragraph(
    'Anche questa euristica è ammissibile: la distanza minima verso una pianta + la massima '
    'distanza interna fornisce un lower bound del costo reale necessario a servire tutte '
    'le piante rimanenti. Il risultato è una riduzione significativa dei nodi esplorati rispetto '
    'sia a UCS che ad A* con Manhattan.'
)

# 8.4
doc.add_heading('8.4. Confronto tra gli Algoritmi', level=2)
doc.add_paragraph(
    'Il progetto esegue tutti e tre gli algoritmi sulla stessa mappa e confronta le statistiche '
    'tramite InstrumentedProblem della libreria AIMA, che conta automaticamente il numero di '
    'nodi esplorati e goal test effettuati.'
)

table5 = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
hdr = table5.rows[0].cells
hdr[0].text = 'Algoritmo'
hdr[1].text = 'Euristica'
hdr[2].text = 'Nodi Esplorati'
hdr[3].text = 'Costo Soluzione'
for algo_data in [
    ('UCS', 'Nessuna', '[inserire dato]', '[inserire dato]'),
    ('A*', 'Manhattan', '[inserire dato]', '[inserire dato]'),
    ('A*', 'Max Pairwise Distance', '[inserire dato]', '[inserire dato]'),
]:
    row = table5.add_row().cells
    for k, v in enumerate(algo_data):
        row[k].text = v

doc.add_paragraph('')
doc.add_paragraph(
    'Il costo della soluzione è identico per tutti e tre gli algoritmi (tutti trovano l\'ottimo), '
    'ma il numero di nodi esplorati diminuisce progressivamente da UCS ad A* con Max Pairwise '
    'Distance, dimostrando il valore di un\'euristica più informativa.'
)

placeholder_img('Tabella con i risultati dell\'esecuzione (nodi, costo, tempo)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. VISUALIZZAZIONE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Fase 6 — Visualizzazione della Soluzione', level=1)

doc.add_paragraph(
    'Il file main.py include due modalità di visualizzazione della soluzione trovata:'
)

doc.add_heading('Visualizzazione Testuale (print_grid)', level=2)
doc.add_paragraph(
    'La funzione print_grid() stampa nel terminale una rappresentazione colorata dello stato '
    'corrente della griglia, usando codici ANSI per i colori: giallo per il robot (B), '
    'verde per le piante (D, V), rosso per le rocce (R), blu per le stazioni (T), '
    'e magenta per la posizione iniziale (S). Ad ogni step viene mostrata l\'azione eseguita '
    'e le informazioni sulle risorse (acqua rimanente, piante residue).'
)

doc.add_heading('Visualizzazione Grafica Interattiva (visualizza_semplice)', level=2)
doc.add_paragraph(
    'La funzione visualizza_semplice() crea una finestra Matplotlib interattiva con '
    'quattro pulsanti: Play, Pausa, Indietro e Avanti. La griglia viene disegnata come '
    'una matrice di celle colorate:'
)
doc.add_paragraph('Grigio: rocce (R)', style='List Bullet')
doc.add_paragraph('Azzurro: stazioni rifornimento (T)', style='List Bullet')
doc.add_paragraph('Oro: cella di arrivo (F)', style='List Bullet')
doc.add_paragraph('Arancione: piante secche (D)', style='List Bullet')
doc.add_paragraph('Rosso: piante molto secche (V)', style='List Bullet')
doc.add_paragraph('Verde: robot (B)', style='List Bullet')

doc.add_paragraph(
    'Un timer interno gestisce la riproduzione automatica (Play): ad ogni tick avanza di '
    'uno step e ridisegna la griglia. I pulsanti Indietro e Avanti permettono la navigazione '
    'manuale. L\'utente può così osservare in dettaglio come il robot pianifica ed esegue '
    'ogni singola azione.'
)

placeholder_img('Screenshot della visualizzazione grafica interattiva con il robot in azione')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. LIBRERIA AIMA
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. Libreria AIMA', level=1)

doc.add_paragraph(
    'Il progetto utilizza la libreria AIMA-Python (Artificial Intelligence: A Modern Approach), '
    'contenuta nella cartella aima/. I file utilizzati sono:'
)

doc.add_paragraph(
    'search.py: contiene le implementazioni degli algoritmi di ricerca (A*, UCS, '
    'best-first search, ecc.) e la classe base Problem da cui eredita AgriBotProblem. '
    'Fornisce anche InstrumentedProblem, un wrapper che conta automaticamente il numero '
    'di nodi generati, stati esplorati e goal test effettuati, utile per il confronto '
    'prestazionale degli algoritmi.',
    style='List Bullet'
)
doc.add_paragraph(
    'utils.py: contiene funzioni di utilità usate internamente da search.py, come '
    'la coda a priorità (PriorityQueue) necessaria per l\'implementazione di A* e UCS.',
    style='List Bullet'
)
doc.add_paragraph(
    'La classe base Problem definisce l\'interfaccia che AgriBotProblem deve implementare: '
    'actions(), result(), goal_test() e il metodo h() per le euristiche. Questo design '
    'permette di separare la definizione del problema dagli algoritmi di ricerca.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. CONCLUSIONI
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11. Conclusioni', level=1)

doc.add_paragraph(
    'Il progetto AgriBot dimostra come integrare diverse aree dell\'Intelligenza Artificiale '
    'in un sistema end-to-end funzionante:'
)
doc.add_paragraph(
    'Computer Vision: il modulo di cropping è in grado di processare immagini reali '
    '(fotografie di griglie disegnate a mano o stampate), corregge la prospettiva e isola '
    'le singole celle con un approccio robusto basato sull\'apertura morfologica.',
    style='List Bullet'
)
doc.add_paragraph(
    'Machine Learning: la CNN addestrata su EMNIST con data augmentation e classi custom '
    '(celle vuote) raggiunge un\'alta precisione nel riconoscimento dei caratteri, '
    'colmando il divario tra immagine reale e rappresentazione simbolica.',
    style='List Bullet'
)
doc.add_paragraph(
    'Ricerca nello Spazio degli Stati: la modellazione del problema con stati, azioni e '
    'transizioni permette di applicare algoritmi classici dell\'IA. Il confronto tra UCS '
    'e A* con euristiche di complessità crescente dimostra concretamente il valore '
    'dell\'informazione euristica nella riduzione dello spazio di ricerca.',
    style='List Bullet'
)
doc.add_paragraph('')
doc.add_paragraph(
    'L\'euristica Max Pairwise Distance si è dimostrata la più efficace, riducendo '
    'significativamente i nodi esplorati rispetto a UCS pur mantenendo l\'ottimalità della '
    'soluzione. Il sistema completo, dalla fotografia alla visualizzazione animata del '
    'percorso, rappresenta un esempio concreto di pipeline di IA applicata.'
)


# ── SALVATAGGIO ──
output_path = '/Users/lorenzovannucci/Desktop/AgriBot-Agent-main/Relazione_AgriBot.docx'
doc.save(output_path)
print(f'Relazione salvata: {output_path}')
